# Word recognition

import pytesseract as pt
import cv2 as cv
import docx
import zipfile
import docx2txt
import re
import docxpy

### Normal jpg pic read ###################################################################
# FIle path : - For image1.png - "C:\Users\sshar127\Desktop\Python"
pt.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
img = cv.imread('image1.png')
txt = pt.image_to_string(img)   


pt.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
img = cv.imread('image1.png')
txt = pt.image_to_string(img)    
# txt = pt.image_to_boxes(img)    # Character and bounding box info
cv.imshow('Result', img)
cv.waitKey(0)

# Detection
img2 = cv.cvtColor(img, cv.COLOR_BGR2RGB)
hImg, wImg = img2.shape

################===============================################################################
pt.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
text = docx2txt.process("Document1.docx", r"C:\Users\sshar127\Desktop\Learning\Office learning DS foundation\Week 9 - NLP")

img = cv.imread(text)
txt = pt.image_to_string(img) 

################===============================################################################
doc = docx.Document("Document1.docx")
result = [p.text for p in doc.paragraphs]

################===============================################################################
document = docx.Document("Document1.docx")

type(document)

document.paragraphs

type(document.paragraphs)

document.paragraphs[1].text

index = 0
for p in document.paragraphs:
    index+=1
    if len(p.text)>0:
        print ('\n Paragraph', index)
        print (p.text)


document = "Document1.docx"
text = docxpy.process(document)

doc = docxpy.DOCReader(document)
doc.process()  # process file
hyperlinks = doc.data['links']


pt.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
doc = docx.Document("Document1.docx")

img = cv.imread(docx2txt.process("Document1.docx", r"C:\Users\sshar127\Desktop\Learning\Office learning DS foundation\Week 9 - NLP"))
img = cv.imread(text)

txt = pt.image_to_string(img) 


doc = docx.Document('Document1.docx')
for i in doc.inline_shapes:
    img = cv.imread(i)
    txt = pt.image_to_string(img)
    print (i.width, i.height)

img = cv.imread('image.png')
txt = pt.image_to_string(img)
















